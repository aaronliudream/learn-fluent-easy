# -*- coding: utf-8 -*-
"""九年级 Unit 1 返修构建:JSON(灌库)+ 同源 Word(审阅)。
语法 3 纯考点×20;阅读 6 篇×5;完形 10 空;听力 10;写作;通关(完形带原文)。
词汇=课本真实词表;其余全部原创,课文只做要点提纲。"""
import json, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DIR = r"C:\Projects\learn-fluent-easy\scripts\g9\u1"
DOCX = r"C:\Users\willi\OneDrive\Desktop\英语教材\初中英语教材\人教版\9年级内容整理\九年级 Unit 1 How can we become good learners.docx"
V, U = "g9", "U1"

# ============================================================ 词汇(真实词表)
VOCAB = {"volume":V,"unit":U,"topic":"How can we become good learners?",
"words":[
 {"word":"textbook","ipa":"/ˈtekstbʊk/","pos":"n.","meaning_cn":"教科书；课本"},
 {"word":"conversation","ipa":"/ˌkɒnvəˈseɪʃn/","pos":"n.","meaning_cn":"交谈；谈话"},
 {"word":"aloud","ipa":"/əˈlaʊd/","pos":"adv.","meaning_cn":"出声地；大声地"},
 {"word":"pronunciation","ipa":"/prəˌnʌnsiˈeɪʃn/","pos":"n.","meaning_cn":"发音；读音"},
 {"word":"sentence","ipa":"/ˈsentəns/","pos":"n.","meaning_cn":"句子"},
 {"word":"patient","ipa":"/ˈpeɪʃnt/","pos":"adj./n.","meaning_cn":"有耐心的；病人"},
 {"word":"expression","ipa":"/ɪkˈspreʃn/","pos":"n.","meaning_cn":"表情；表达方式"},
 {"word":"discover","ipa":"/dɪˈskʌvə(r)/","pos":"v.","meaning_cn":"发现；发觉"},
 {"word":"secret","ipa":"/ˈsiːkrət/","pos":"n./adj.","meaning_cn":"秘密；秘诀；秘密的"},
 {"word":"grammar","ipa":"/ˈɡræmə(r)/","pos":"n.","meaning_cn":"语法"},
 {"word":"repeat","ipa":"/rɪˈpiːt/","pos":"v.","meaning_cn":"重复；重做"},
 {"word":"note","ipa":"/nəʊt/","pos":"n./v.","meaning_cn":"笔记；记录；注意"},
 {"word":"pronounce","ipa":"/prəˈnaʊns/","pos":"v.","meaning_cn":"发音"},
 {"word":"increase","ipa":"/ɪnˈkriːs/","pos":"v.","meaning_cn":"增加；增长"},
 {"word":"partner","ipa":"/ˈpɑːtnə(r)/","pos":"n.","meaning_cn":"搭档；同伴"},
 {"word":"physics","ipa":"/ˈfɪzɪks/","pos":"n.","meaning_cn":"物理（学）"},
 {"word":"chemistry","ipa":"/ˈkemɪstri/","pos":"n.","meaning_cn":"化学"},
 {"word":"born","ipa":"/bɔːn/","pos":"adj.","meaning_cn":"天生的（be born 出生）"},
 {"word":"create","ipa":"/kriˈeɪt/","pos":"v.","meaning_cn":"创造；创建"},
 {"word":"speed","ipa":"/spiːd/","pos":"n.","meaning_cn":"速度"},
 {"word":"ability","ipa":"/əˈbɪləti/","pos":"n.","meaning_cn":"能力；才能"},
 {"word":"active","ipa":"/ˈæktɪv/","pos":"adj.","meaning_cn":"积极的；活跃的"},
 {"word":"brain","ipa":"/breɪn/","pos":"n.","meaning_cn":"大脑"},
 {"word":"connect","ipa":"/kəˈnekt/","pos":"v.","meaning_cn":"连接；联系"},
 {"word":"review","ipa":"/rɪˈvjuː/","pos":"v./n.","meaning_cn":"复习；回顾"},
 {"word":"knowledge","ipa":"/ˈnɒlɪdʒ/","pos":"n.","meaning_cn":"知识；学问"},
 {"word":"wisely","ipa":"/ˈwaɪzli/","pos":"adv.","meaning_cn":"明智地"},
 {"word":"memorize","ipa":"/ˈmeməraɪz/","pos":"v.","meaning_cn":"记忆；记住"},
 {"word":"attention","ipa":"/əˈtenʃn/","pos":"n.","meaning_cn":"注意；关注"},
],
"phrases":[
 {"phrase":"pay attention to","meaning_cn":"注意；关注"},
 {"phrase":"connect ... with ...","meaning_cn":"把……和……连接（联系）起来"},
 {"phrase":"read aloud","meaning_cn":"朗读；大声读"},
 {"phrase":"make word cards","meaning_cn":"制作单词卡片"},
 {"phrase":"be interested in","meaning_cn":"对……感兴趣"},
 {"phrase":"be good at","meaning_cn":"擅长（做）……"},
 {"phrase":"fall in love with","meaning_cn":"爱上；喜欢上"},
 {"phrase":"the more ..., the more ...","meaning_cn":"越……，就越……"},
 {"phrase":"look up","meaning_cn":"（在词典里）查（单词）"},
 {"phrase":"for example","meaning_cn":"例如"},
]}

# ============================================================ 语法 3 考点 × 20
P1="动词 + by + 动名词(by doing)表示“通过某种方式”"
P2="How 引导的方法疑问句(询问做事方式,答语为方式)"
P3="情态动词/助动词 + 动词原形"
C1,C2,C3="g9u1.01","g9u1.02","g9u1.03"
def gq(qid,point,code,stem,options,ai,expl):
    return {"qid":qid,"volume":V,"unit":U,"point":point,"code":code,"stem":stem,
            "options":options,"answer_index":ai,"answer_text":options[ai],"explanation":expl}

G1=[
 ("q1","— How do you improve your English? — I improve it ____ English songs.",["by sing","by singing","to sing","singing"],1,"by 后接动名词。"),
 ("q2","We can learn a lot ____ books about history.",["by reading","to reading","by read","read"],0,"by + 动名词 reading。"),
 ("q3","My sister became a better writer ____ every day.",["by write","writes","write","by writing"],3,"by + 动名词 writing。"),
 ("q4","— How did you remember so many words? — ____ them again and again.",["By repeating","Repeat","By repeat","To repeat"],0,"句首方式状语 By + doing。"),
 ("q5","Tom keeps fit ____ football after school.",["playing","by play","to play","by playing"],3,"by + 动名词 playing。"),
 ("q6","She learns new words ____ word cards.",["by make","making","by making","to make"],2,"by + 动名词 making。"),
 ("q7","You can improve your pronunciation ____ to tapes.",["listen","by listening","listening","by listen"],1,"by + 动名词 listening。"),
 ("q8","— How does he study for tests? — ____ with a study group.",["Working","By work","By working","To work"],2,"句首方式状语 By + doing。"),
 ("q9","I relax ____ music after class.",["by listen to","listening to","to listen to","by listening to"],3,"by + 动名词短语 listening to。"),
 ("q10","He saved money ____ part-time jobs.",["by doing","by do","doing","to do"],0,"by + 动名词 doing。"),
 ("q11","We protect our eyes ____ less TV.",["by watch","watching","by watching","watch"],2,"by + 动名词 watching。"),
 ("q12","— How did you learn to swim? — ____ every weekend.",["By practising","Practise","By practise","To practise"],0,"句首 By + doing。"),
 ("q13","She got good grades ____ hard.",["by study","studying","to study","by studying"],3,"by + 动名词 studying。"),
 ("q14","You can make friends ____ in clubs.",["by join","by joining","joining","to join"],1,"by + 动名词 joining。"),
 ("q15","They celebrated ____ a big party.",["by have","having","by having","to have"],2,"by + 动名词 having。"),
 ("q16","— How do you train your memory? — ____ poems.",["By reciting","Recite","By recite","To recite"],0,"句首 By + doing。"),
 ("q17","He improved his handwriting ____ slowly and carefully.",["by write","writing","by writing","to write"],2,"by + 动名词 writing。"),
 ("q18","We learn about the world ____ to new places.",["by travel","by travelling","travelling","to travel"],1,"by + 动名词 travelling。"),
 ("q19","She keeps healthy ____ vegetables every day.",["by eat","eating","to eat","by eating"],3,"by + 动名词 eating。"),
 ("q20","— How can I calm down? — ____ deeply.",["By breathing","Breathe","By breathe","To breathe"],0,"句首 By + doing。"),
]
G2=[
 ("q1","— ____ you study for a test? — I study by making notes.",["What do","How do","How are","Why do"],1,"问方式用 How do you...?;答语为 by + doing。"),
 ("q2","— ____ can I improve my spoken English quickly? — By talking more.",["How","What","Where","Who"],0,"问“怎样(方法)”用 How。"),
 ("q3","— ____ did you find the meaning of the word? — By guessing from the sentence.",["When","What","How","Which"],2,"答语为方式,故用 How。"),
 ("q4","— ____ do you go to school? — By bus.",["What","Where","What time","How"],3,"问交通方式用 How,答 By bus。"),
 ("q5","— ____ do you keep healthy? — By exercising every day.",["What","How","When","Who"],1,"答语为方式,用 How。"),
 ("q6","— ____ can we become good learners? — By studying actively.",["What","Why","How","Where"],2,"答语为方式,用 How。"),
 ("q7","— ____ did they win the game? — By working together.",["How","When","What","Who"],0,"答语为方式,用 How。"),
 ("q8","— ____ does Lily learn so well? — By reviewing her notes daily.",["Where","Why","What","How"],3,"答语为方式,用 How。"),
 ("q9","— ____ can I remember new words better? — By using them in sentences.",["How","What","When","Which"],0,"答语为方式,用 How。"),
 ("q10","— ____ do you usually relax? — By listening to music.",["Why","Where","What","How"],3,"答语为方式,用 How。"),
 ("q11","— ____ do you save water at home? — By turning off the tap.",["What","How","Who","Where"],1,"答语为方式,用 How。"),
 ("q12","— ____ did you fix the computer? — By following the instructions.",["How","When","Why","What"],0,"答语为方式,用 How。"),
 ("q13","— ____ can students reduce stress? — By planning their time wisely.",["What","Why","How","Who"],2,"答语为方式,用 How。"),
 ("q14","— ____ do farmers protect the crops? — By building fences.",["When","How","What","Where"],1,"答语为方式,用 How。"),
 ("q15","— ____ did she pass the exam? — By studying hard for a month.",["Why","What","How","Who"],2,"答语为方式,用 How。"),
 ("q16","— ____ can I improve my writing? — By reading good articles.",["How","What","Where","When"],0,"答语为方式,用 How。"),
 ("q17","— ____ do you learn idioms? — By watching English films.",["What","Why","Who","How"],3,"答语为方式,用 How。"),
 ("q18","— ____ did the team raise money? — By selling old books.",["When","How","What","Where"],1,"答语为方式,用 How。"),
 ("q19","— ____ can we learn about other cultures? — By travelling and reading.",["What","Why","How","Which"],2,"答语为方式,用 How。"),
 ("q20","— ____ do you train your dog? — By giving it small rewards.",["When","What","Where","How"],3,"答语为方式,用 How。"),
]
G3=[
 ("q1","— How ____ Lily learn English so well? — By reading aloud.",["do","is","did","does"],3,"主语三单、一般现在时,助动词 does + 动词原形。"),
 ("q2","— How can I ____ my listening? — By listening to the radio.",["improving","improve","improves","improved"],1,"can 后接动词原形。"),
 ("q3","— How do you ____ new words? — I memorize them by writing.",["remember","remembers","remembering","to remember"],0,"do you 后接动词原形。"),
 ("q4","— How can we ____ good learners? — By studying actively.",["becoming","became","to become","become"],3,"can 后接动词原形。"),
 ("q5","You should ____ more water every day.",["drink","drinks","drinking","drank"],0,"should 后接动词原形。"),
 ("q6","We must ____ attention to road safety.",["pays","paying","pay","paid"],2,"must 后接动词原形。"),
 ("q7","Students should not ____ in class.",["talks","talk","talking","talked"],1,"should not 后接动词原形。"),
 ("q8","____ he speak three languages?",["Does","Do","Is","Are"],0,"主语 he 三单,助动词 Does + 原形。"),
 ("q9","They didn't ____ the meeting yesterday.",["attended","attends","attending","attend"],3,"didn't 后接动词原形。"),
 ("q10","May I ____ your dictionary for a minute?",["borrowing","borrows","borrow","borrowed"],2,"May 后接动词原形。"),
 ("q11","____ you finish your homework last night?",["Do","Did","Does","Are"],1,"一般过去时疑问,助动词 Did + 原形。"),
 ("q12","She can ____ the piano very well.",["play","plays","playing","played"],0,"can 后接动词原形。"),
 ("q13","We will ____ our best to win.",["trying","try","tries","tried"],1,"will 后接动词原形。"),
 ("q14","____ your parents like Beijing opera?",["Does","Is","Do","Was"],2,"主语复数,助动词 Do + 原形。"),
 ("q15","He doesn't ____ coffee in the evening.",["drinks","drinking","drank","drink"],3,"doesn't 后接动词原形。"),
 ("q16","Could you ____ me the way to the station?",["tell","telling","tells","told"],0,"Could 后接动词原形。"),
 ("q17","Tom must ____ his room before going out.",["cleans","clean","cleaning","cleaned"],1,"must 后接动词原形。"),
 ("q18","____ the students go on a trip last week?",["Do","Does","Did","Are"],2,"一般过去时疑问,助动词 Did + 原形。"),
 ("q19","You'd better ____ to bed early tonight.",["going","goes","went","go"],3,"had better 后接动词原形。"),
 ("q20","We should ____ down our ideas in a notebook.",["writing","writes","write","wrote"],2,"should 后接动词原形。"),
]
def build_points(rows,point,code):
    return [gq(f"{code}.{r[0]}",point,code,r[1],r[2],r[3],r[4]) for r in rows]
GRAMMAR={"volume":V,"unit":U,"points":[
 {"code":C1,"point":P1,"overview":"介词 by 后接动名词(v-ing),作方式状语,表示“通过做某事”,常回答 How 问句。by 后用 doing,不用动词原形;与表交通方式的 by bus 区分。"},
 {"code":C2,"point":P2,"overview":"询问“怎样做某事/方式”用 How 引导的特殊疑问句,答语常用 by + doing。判断靠“答语是方式”而非后面的动词形式。"},
 {"code":C3,"point":P3,"overview":"情态动词(can/should/must/may/could/will/had better 等)和助动词(do/does/did)后一律接动词原形。"},
], "questions": build_points(G1,P1,C1)+build_points(G2,P2,C2)+build_points(G3,P3,C3)}

# ============================================================ 阅读 6 篇 × 5
def rq(qid,code,stem,options,ai,expl):
    return {"qid":qid,"volume":V,"unit":U,"code":code,"stem":stem,"options":options,
            "answer_index":ai,"answer_text":options[ai],"explanation":expl}
RD_PASSAGES=[
 {"code":"g9u1.rd1","title":"Becoming a Good Learner","genre":"说明文",
  "body":("Good learners are not born; they are made. They believe that everyone can learn well by "
   "finding the right way. First, good learners are active. They do not just wait for the teacher to "
   "tell them everything. Instead, they ask questions and try to find the answers by themselves. "
   "Second, good learners are patient. They know that making mistakes is a normal part of learning, "
   "so they never give up. Third, good learners connect new knowledge with what they already know. "
   "For example, when they meet a new word, they make a sentence with it so that they can remember "
   "it better. Finally, good learners study wisely. They plan their time and review what they have "
   "learned often. If you follow these tips, you can become a good learner, too."),
  "q":[
   ("q1","What is the passage mainly about?",["How to become a good learner.","The history of English.","How to make friends.","Why exams are hard."],0,"全文讲怎样成为优秀学习者。"),
   ("q2","Good learners are active because they ____.",["wait for the teacher","ask questions and find answers themselves","never study at home","only read textbooks"],1,"原文:they ask questions and try to find the answers by themselves。"),
   ("q3","What do good learners think about mistakes?",["They are terrible.","They should be hidden.","They are a normal part of learning.","They mean failure."],2,"原文:making mistakes is a normal part of learning。"),
   ("q4","How do good learners remember a new word?",["By reading it once.","By writing it many times.","By ignoring it.","By making a sentence with it."],3,"原文:make a sentence with it so that they can remember it better。"),
   ("q5","The word “wisely” here means ____.",["in a lazy way","in a smart and sensible way","in a hurry","by luck"],1,"wisely=明智地。"),
  ]},
 {"code":"g9u1.rd2","title":"Anna and Her Reading Journey","genre":"记叙文",
  "body":("When Anna was young, she hated reading. The books at school looked boring, and she could not "
   "understand many words. One summer, her aunt gave her a book about a girl who travelled around the "
   "world. Anna started reading just one page a day. Slowly, she wanted to know what would happen "
   "next, so she read more and more. Whenever she met a new word, she did not stop; she guessed its "
   "meaning from the story. By the end of the summer, she had finished five books. Reading had become "
   "her favourite hobby. Now Anna says that the secret is to start with something you really enjoy."),
  "q":[
   ("q1","How did Anna feel about reading when she was young?",["She loved it.","She found it easy.","She hated it.","She never heard of it."],2,"原文:she hated reading。"),
   ("q2","Who gave Anna the book?",["Her mother.","Her teacher.","Her friend.","Her aunt."],3,"原文:her aunt gave her a book。"),
   ("q3","What did Anna do when she met a new word?",["She guessed its meaning from the story.","She looked it up at once.","She skipped the whole book.","She asked her aunt."],0,"原文:she guessed its meaning from the story。"),
   ("q4","How many books did Anna finish by the end of summer?",["Three.","Five.","Two.","Ten."],1,"原文:she had finished five books。"),
   ("q5","What is Anna's secret to enjoying reading?",["Read very fast.","Use a dictionary often.","Start with something you enjoy.","Read one book a year."],2,"原文:start with something you really enjoy。"),
  ]},
 {"code":"g9u1.rd3","title":"Two Ways to Remember Words","genre":"对比说明",
  "body":("Students often ask how to remember English words. There are two common ways. The first way is "
   "to write a word many times. This helps some students, but it can be boring, and you may forget the "
   "word soon. The second way is to use the word in real sentences. For example, if you learn the word "
   "“brave”, you can say, “My brother is brave.” This way takes more time at first, but you will "
   "remember the word much longer, because it is connected with a real idea. Many teachers believe the "
   "second way is better, because learning with meaning is stronger than learning by repeating alone."),
  "q":[
   ("q1","What is the passage mainly about?",["The history of English words.","How to write quickly.","Why words are difficult.","Two ways to remember words."],3,"主旨:记单词的两种方法。"),
   ("q2","What is a problem with writing a word many times?",["It can be boring and easily forgotten.","It costs a lot of money.","It needs a partner.","It is not allowed."],0,"原文:it can be boring, and you may forget the word soon。"),
   ("q3","The example “My brother is brave” shows ____.",["writing a word many times","using a word in a real sentence","reading a word aloud","drawing a picture"],1,"该例说明“在真实句子中使用单词”。"),
   ("q4","Why is the second way better?",["It is faster.","It is easier.","Learning with meaning is stronger.","It needs no time."],2,"原文:learning with meaning is stronger。"),
   ("q5","Which way do many teachers prefer?",["The first way.","Neither way.","Both equally.","The second way."],3,"原文:the second way is better。"),
  ]},
 {"code":"g9u1.rd4","title":"A Letter to a New Classmate","genre":"书信",
  "body":("Dear Sam, Welcome to our class! I know learning English in a new school can be hard, so here "
   "are some tips. First, do not be afraid to make mistakes when you speak. Everyone makes them, and "
   "they help us learn. Second, try to read English for a few minutes every day, even just a comic "
   "book. Third, find a study partner so that you can practise together. Finally, be patient with "
   "yourself, because real progress takes time. If you ever need help, just ask me. I am sure you "
   "will do well. Your friend, Lily"),
  "q":[
   ("q1","Why did Lily write this letter?",["To welcome Sam and give him study tips.","To invite Sam to a party.","To say goodbye to Sam.","To borrow a book."],0,"信的目的是欢迎并给学习建议。"),
   ("q2","What does Lily say about mistakes?",["They are useless.","They help us learn.","They should be hidden.","They are dangerous."],1,"原文:they help us learn。"),
   ("q3","How often should Sam read English?",["Once a week.","Only at weekends.","A few minutes every day.","Three hours a day."],2,"原文:read English for a few minutes every day。"),
   ("q4","Why should Sam find a study partner?",["To copy answers.","To make money.","To save time.","To practise together."],3,"原文:so that you can practise together。"),
   ("q5","What is Lily's last piece of advice?",["Be patient with yourself.","Buy many books.","Study all night.","Never ask for help."],0,"原文:be patient with yourself。"),
  ]},
 {"code":"g9u1.rd5","title":"The Power of Mistakes","genre":"议论文",
  "body":("Many students are afraid of making mistakes. They keep quiet in class because they do not want "
   "to look silly. However, mistakes are not bad at all. In fact, they are an important part of "
   "learning. When you make a mistake and then correct it, your brain remembers the right answer "
   "better. Think about babies: they fall down many times before they learn to walk. Learning a "
   "language is the same. Every mistake shows you what to study next. So, the next time you make a "
   "mistake, do not feel sad. Be brave, learn from it, and try again. That is how good learners grow."),
  "q":[
   ("q1","Why do some students keep quiet in class?",["They are tired.","They are afraid of looking silly.","They have no questions.","They dislike English."],1,"原文:they do not want to look silly。"),
   ("q2","What happens when you correct a mistake?",["Nothing happens.","You feel sad.","Your brain remembers the right answer better.","You forget everything."],2,"原文:your brain remembers the right answer better。"),
   ("q3","Why does the writer mention babies?",["To talk about families.","To make readers laugh.","To change the topic.","To show that mistakes are part of learning."],3,"以婴儿学步类比,说明犯错是学习的一部分。"),
   ("q4","What does every mistake show you?",["What to study next.","Which friend to choose.","When to sleep.","Where to go."],0,"原文:every mistake shows you what to study next。"),
   ("q5","What does the writer advise readers to do?",["Give up quickly.","Be brave and learn from mistakes.","Stay silent.","Avoid all tests."],1,"原文:Be brave, learn from it, and try again。"),
  ]},
 {"code":"g9u1.rd6","title":"Learning Beyond the Classroom","genre":"应用说明",
  "body":("School is not the only place to learn English. There are many fun ways to practise outside the "
   "classroom. You can watch English cartoons or films with subtitles to train your ears. You can "
   "listen to English songs and sing along to improve your pronunciation. The Internet is also "
   "helpful: you can read short news stories or chat with online friends from other countries. Some "
   "students keep a diary in English to practise writing. The most important thing is to use English "
   "a little every day. When learning becomes part of your daily life, you will improve without even "
   "noticing it."),
  "q":[
   ("q1","What is the passage mainly about?",["The rules of grammar.","A famous English teacher.","Ways to practise English outside class.","The history of films."],2,"主旨:课外练英语的方法。"),
   ("q2","How can watching films help you?",["They cost nothing.","They are short.","They are funny.","They train your ears."],3,"原文:watch films with subtitles to train your ears。"),
   ("q3","Why listen to English songs?",["To improve pronunciation.","To learn grammar rules.","To make friends.","To pass exams."],0,"原文:sing along to improve your pronunciation。"),
   ("q4","What do some students do to practise writing?",["Watch more TV.","Keep a diary in English.","Play games.","Sing songs."],1,"原文:keep a diary in English to practise writing。"),
   ("q5","What is the most important thing?",["Study only before exams.","Buy expensive books.","Use English a little every day.","Travel abroad."],2,"原文:use English a little every day。"),
  ]},
]
READING={"volume":V,"unit":U,"passages":[
 {"code":p["code"],"title":p["title"],"genre":p["genre"],"body":p["body"],
  "questions":[rq(f'{p["code"]}.{r[0]}',p["code"],r[1],r[2],r[3],r[4]) for r in p["q"]]}
 for p in RD_PASSAGES]}

# ============================================================ 完形 10 空
CLOZE_TEXT=("Li Lei used to be weak in English. He could not understand the teacher and was afraid to "
"__1__ in class. Then he decided to change. Every morning he read English __2__ for twenty minutes to "
"improve his pronunciation. He also built his __3__ by making word cards and reviewing them on the "
"bus. When he met a new word, he tried to guess its meaning by __4__ the sentences around it. In "
"class, he became more __5__ and often asked the teacher questions. He knew that everyone makes "
"__6__, so he was __7__ with himself and never gave up. He also joined an English club, where he made "
"many __8__. Little by little, his English __9__. Now he can have a conversation with foreign "
"visitors easily. Li Lei says that anyone can learn English well by working hard and finding the "
"right __10__.")
def cq(qid,n,options,ai,expl):
    return {"qid":qid,"volume":V,"unit":U,"code":"g9u1.cz","blank":n,"options":options,
            "answer_index":ai,"answer_text":options[ai],"explanation":expl}
CLOZE={"volume":V,"unit":U,"code":"g9u1.cz","text":CLOZE_TEXT,"questions":[
 cq("g9u1.cz.q1",1,["sleep","speak","sing","run"],1,"afraid to speak in class 害怕在课上发言。"),
 cq("g9u1.cz.q2",2,["aloud","quietly","slowly","alone"],0,"read aloud 朗读以改善发音。"),
 cq("g9u1.cz.q3",3,["grammar","vocabulary","pronunciation","knowledge"],1,"make word cards 是为扩大 vocabulary。"),
 cq("g9u1.cz.q4",4,["writing","copying","reading","drawing"],2,"by reading the sentences around it 读上下文猜词义。"),
 cq("g9u1.cz.q5",5,["quiet","shy","tired","active"],3,"常提问→更 active 积极。"),
 cq("g9u1.cz.q6",6,["mistakes","friends","notes","plans"],0,"everyone makes mistakes 人人都会犯错。"),
 cq("g9u1.cz.q7",7,["angry","patient","busy","strict"],1,"对自己有耐心 patient,不放弃。"),
 cq("g9u1.cz.q8",8,["problems","mistakes","friends","enemies"],2,"加入俱乐部→结交许多 friends。"),
 cq("g9u1.cz.q9",9,["fell","stopped","improved","disappeared"],2,"坚持努力→英语 improved 提高。"),
 cq("g9u1.cz.q10",10,["time","place","reason","way"],3,"find the right way 找到正确方法,呼应全文。"),
]}

# ============================================================ 听力(原创对话+短文+10题)
DIALOGUE=["Lin Tao: Hi, Sarah! The English exam is coming. How do you study for it?",
 "Sarah: I study by reading English stories every night. It helps me learn new words.",
 "Lin Tao: That's a good idea. But I always forget new words quickly.",
 "Sarah: You can try making word cards and reviewing them on the bus.",
 "Lin Tao: Sounds useful. What about listening? Mine is really weak.",
 "Sarah: Don't worry. I improved my listening by watching short English videos.",
 "Lin Tao: How long do you watch them every day?",
 "Sarah: About twenty minutes. Be patient, and you'll get better step by step.",
 "Lin Tao: Thank you, Sarah. I'll start tonight!"]
DIALOGUE_CN=("林涛:嗨,莎拉!英语考试快到了,你怎么备考?/莎拉:我每晚通过读英语故事来学习,这能帮我记新词。/"
"林涛:好主意。可我总是很快忘掉新词。/莎拉:你可以试试做单词卡片,在公交车上复习。/林涛:听起来有用。听力呢?我很弱。/"
"莎拉:别担心,我是通过看英语短视频提高听力的。/林涛:你每天看多久?/莎拉:大约二十分钟。要有耐心,你会一步步进步。/"
"林涛:谢谢你,莎拉!我今晚就开始!")
PASSAGE=("My name is Mike. Last year my English was poor, but now I am one of the best students in my "
"class. How did I change? First, I learned grammar by doing exercises every day. Second, I increased "
"my vocabulary by reading newspapers and writing down useful expressions. I also practised speaking "
"by talking with my foreign teacher after class. At first it was hard, but I never gave up. Now I "
"really enjoy learning English, and I believe everyone can become a good learner by working hard and "
"finding the right way.")
PASSAGE_CN=("我叫迈克。去年我的英语很差,但现在我是班上最好的学生之一。我怎么改变的?第一,我每天通过做练习学语法。"
"第二,我通过读报纸、记下有用表达来扩大词汇量。我还通过课后和外教交谈练口语。起初很难,但我从不放弃。"
"现在我很享受学英语,我相信只要努力并找到正确方法,人人都能成为优秀的学习者。")
def lq(qid,stem,options,ai):
    return {"qid":qid,"volume":V,"unit":U,"code":"g9u1.ls","stem":stem,"options":options,
            "answer_index":ai,"answer_text":options[ai]}
LISTENING={"volume":V,"unit":U,"code":"g9u1.ls","materials":[
 {"type":"dialogue","transcript":DIALOGUE,"translation_cn":DIALOGUE_CN},
 {"type":"passage","transcript":[PASSAGE],"translation_cn":PASSAGE_CN}],
 "questions":[
 lq("g9u1.ls.q1","How does Sarah study for the English exam?",["By playing games.","By singing songs.","By reading English stories.","By drawing pictures."],2),
 lq("g9u1.ls.q2","What does Sarah suggest for remembering words?",["Making word cards.","Reading aloud.","Copying the textbook.","Asking the teacher."],0),
 lq("g9u1.ls.q3","How did Sarah improve her listening?",["By taking notes.","By reading novels.","By making friends.","By watching short English videos."],3),
 lq("g9u1.ls.q4","How long does Sarah watch videos every day?",["About ten minutes.","About twenty minutes.","About half an hour.","About one hour."],1),
 lq("g9u1.ls.q5","What does Sarah tell Lin Tao to be?",["Careful.","Quiet.","Patient.","Active."],2),
 lq("g9u1.ls.q6","How was Mike's English last year?",["Poor.","Excellent.","Just OK.","The best."],0),
 lq("g9u1.ls.q7","How did Mike learn grammar?",["By listening to tapes.","By watching TV.","By chatting online.","By doing exercises every day."],3),
 lq("g9u1.ls.q8","How did Mike increase his vocabulary?",["By playing word games.","By reading newspapers and writing down expressions.","By memorizing the dictionary.","By writing letters."],1),
 lq("g9u1.ls.q9","How did Mike practise speaking?",["By reading aloud at home.","By recording himself.","By talking with his foreign teacher.","By singing songs."],2),
 lq("g9u1.ls.q10","What does Mike believe?",["Everyone can become a good learner by working hard.","English is useless.","Only clever students learn well.","Learning English is impossible."],0),
]}

# ============================================================ 写作
WRITING={"volume":V,"unit":U,"topic":"How I Study English(我怎样学习英语)",
 "genre":"个人短文 / 给学弟学妹的学习建议",
 "points":["至少写出 3 种你的学习方法","必须使用本单元句型 “… by doing …”","结尾给出一句鼓励或建议"],
 "word_count":"80 词左右",
 "model_essay":("Everyone wants to learn English well, and here is how I do it. I build my vocabulary "
 "by making word cards and reviewing them every morning. I improve my listening by watching English "
 "cartoons, and I practise speaking by talking with my classmates in English. When I meet a new word, "
 "I find out its meaning by reading the sentences around it. Learning a language takes time, so be "
 "patient. If you keep working hard and find the right way, you will surely become a good learner!"),
 "model_translation":("人人都想学好英语,下面是我的做法。我通过制作单词卡片并每天早晨复习来积累词汇。我通过看英语"
 "动画片提高听力,通过和同学用英语交谈来练习口语。遇到生词时,我通过读上下文的句子来弄清词义。学习语言需要时间,"
 "所以要有耐心。只要坚持努力并找到正确的方法,你一定会成为一名优秀的学习者!")}

READING_OUTLINE=[
 "Section A 主线:同学们围绕“如何备考、如何学英语”交流方法——小组合作、朗读练发音、看视频练听力、读词组提高阅读速度;并建议先抓大意、读上下文猜词义,强调要有耐心。",
 "Section B 主线:一名学生分享学英语经历——起初害怕、读不懂,后来通过看电影记单词、加入俱乐部多开口,逐渐爱上英语;说明找到适合自己的方法 + 坚持,人人都能学好。",
 "单元核心:好的学习者 = 主动 + 会用方法(by doing)+ 有耐心 + 善于发现规律。",
]

# ============================================================ 单元通关(完形带原文)
FINALTEST={"volume":V,"unit":U,"code":"g9u1.final","title":"Unit 1 单元通关",
 "note":"从语法/阅读/完形/听力各关已出题目抽题组卷,不另出新题,仅引用 qid。完形题携带其所属短文原文。",
 "cloze_context":CLOZE_TEXT,
 "question_refs":["g9u1.01.q1","g9u1.01.q12","g9u1.02.q2","g9u1.02.q10","g9u1.03.q5","g9u1.03.q16",
   "g9u1.rd1.q1","g9u1.rd3.q2","g9u1.rd5.q4",
   "g9u1.cz.q2","g9u1.cz.q6","g9u1.cz.q9",
   "g9u1.ls.q1","g9u1.ls.q6","g9u1.ls.q9"]}

# ============================================================ 写 JSON
os.makedirs(DIR,exist_ok=True)
def dump(n,o):
    with open(os.path.join(DIR,n),"w",encoding="utf-8") as f: json.dump(o,f,ensure_ascii=False,indent=2)
dump("g9-u1-vocab.json",VOCAB); dump("g9-u1-grammar.json",GRAMMAR); dump("g9-u1-reading.json",READING)
dump("g9-u1-cloze.json",CLOZE); dump("g9-u1-listening.json",LISTENING); dump("g9-u1-writing.json",WRITING)
dump("g9-u1-finaltest.json",FINALTEST)

# ============================================================ Word(同源)
doc=Document(); doc.styles['Normal'].font.name='Microsoft YaHei'; doc.styles['Normal'].font.size=Pt(10.5)
NAVY=RGBColor(0x1c,0x33,0x5c); ORANGE=RGBColor(0xC0,0x5A,0x1B); GREY=RGBColor(0x80,0x80,0x80); LET="ABCD"
def h1(t):
    p=doc.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(16); r.bold=True
def h2(t):
    p=doc.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=ORANGE; r.font.size=Pt(13); r.bold=True
def para(t,bold=False,size=10.5,color=None,italic=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.font.size=Pt(size); r.italic=italic
    if color: r.font.color.rgb=color
def bullet(t): doc.add_paragraph(t,style='List Bullet')
def render_q(qs,show_expl=True):
    for i,q in enumerate(qs,1):
        para(f"{i}. {q['stem']}")
        doc.add_paragraph("   "+"    ".join(f"{LET[j]}. {o}" for j,o in enumerate(q['options'])))
    para("【答案与解析】" if show_expl else "【答案】",bold=True,size=10,color=ORANGE)
    for i,q in enumerate(qs,1):
        doc.add_paragraph(f"{i}. {LET[q['answer_index']]}"+(f" —— {q.get('explanation','')}" if show_expl and q.get('explanation') else ""))

t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=t.add_run("人教版 英语 · 九年级全一册"); r.font.size=Pt(13); r.font.color.rgb=NAVY
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=t.add_run("Unit 1  How can we become good learners?"); r.font.size=Pt(20); r.bold=True; r.font.color.rgb=NAVY
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=t.add_run("我们怎样成为优秀的学习者?"); r.font.size=Pt(12); r.font.color.rgb=ORANGE
para(f"(volume={V} · unit={U} · code 前缀 g9u1)",italic=True,size=9,color=GREY); doc.add_paragraph()

h1("一、单元词汇表 Vocabulary")
para("词汇取自人教版九年级 Unit 1 真实词表;音标为标准 IPA。",italic=True,size=9)
tb=doc.add_table(rows=1,cols=4); tb.style='Light List Accent 1'
for j,x in enumerate(["单词","音标","词性","中文释义"]): tb.rows[0].cells[j].text=x; tb.rows[0].cells[j].paragraphs[0].runs[0].font.bold=True
for w in VOCAB["words"]:
    c=tb.add_row().cells; c[0].text=w["word"]; c[1].text=w["ipa"]; c[2].text=w["pos"]; c[3].text=w["meaning_cn"]; c[0].paragraphs[0].runs[0].font.bold=True
para(""); para("【常用短语 Phrases】",bold=True)
for p in VOCAB["phrases"]: bullet(f'{p["phrase"]} — {p["meaning_cn"]}')

h1("二、语法考点与专项练习 Grammar(3 考点 × 20 题)")
for pt in GRAMMAR["points"]:
    h2(f'考点 {pt["code"][-2:]}：{pt["point"]}')
    para("讲解:"+pt["overview"])
    qs=[q for q in GRAMMAR["questions"] if q["code"]==pt["code"]]
    para(f'◆ 专项练习({len(qs)} 题)',bold=True,size=11,color=NAVY); render_q(qs)

h1("三、听力训练 Listening(原创材料)")
h2("材料 A：对话 Dialogue")
for ln in LISTENING["materials"][0]["transcript"]: doc.add_paragraph(ln)
para("【参考译文】"+LISTENING["materials"][0]["translation_cn"],size=9,italic=True,color=GREY)
h2("材料 B：短文 Passage")
for ln in LISTENING["materials"][1]["transcript"]: doc.add_paragraph(ln)
para("【参考译文】"+LISTENING["materials"][1]["translation_cn"],size=9,italic=True,color=GREY)
para(f'【听力题目({len(LISTENING["questions"])} 题)】',bold=True,size=11,color=NAVY)
render_q(LISTENING["questions"],show_expl=False)

h1("四、完形填空 Cloze(原创短文 · 10 空)")
doc.add_paragraph(CLOZE["text"])
para(f'【完形题目({len(CLOZE["questions"])} 题)】',bold=True,size=11,color=NAVY)
for q in CLOZE["questions"]:
    para(f'{q["blank"]}. '+"    ".join(f"{LET[j]}. {o}" for j,o in enumerate(q["options"])))
para("【答案与解析】",bold=True,size=10,color=ORANGE)
for q in CLOZE["questions"]: doc.add_paragraph(f'{q["blank"]}. {LET[q["answer_index"]]} —— {q["explanation"]}')

h1("五、阅读理解 Reading(原创短文 6 篇)")
for p in READING["passages"]:
    h2(f'{p["title"]}  ({p["genre"]})')
    doc.add_paragraph(p["body"])
    para(f'【阅读题目({len(p["questions"])} 题)】',bold=True,size=10,color=NAVY); render_q(p["questions"])

h1("六、写作 Writing")
h2("写作要求"); bullet("题目:"+WRITING["topic"]); bullet("体裁:"+WRITING["genre"])
for p in WRITING["points"]: bullet(p)
bullet("词数:"+WRITING["word_count"])
h2("范文(原创参考)"); doc.add_paragraph(WRITING["model_essay"])
h2("范文译文"); doc.add_paragraph(WRITING["model_translation"])

h1("七、课文要点(内容概述,非课文原文)")
for o in READING_OUTLINE: bullet(o)

h1("八、单元通关 Final Test(抽题组卷)")
para(FINALTEST["note"],italic=True,size=9)
allq={}
for q in GRAMMAR["questions"]: allq[q["qid"]]=q
for p in READING["passages"]:
    for q in p["questions"]: allq[q["qid"]]=q
for q in CLOZE["questions"]: allq[q["qid"]]=q
for q in LISTENING["questions"]: allq[q["qid"]]=q
picked=[allq[r] for r in FINALTEST["question_refs"]]
if any(q["code"]=="g9u1.cz" for q in picked):
    para("【完形题所属短文】",bold=True,size=10,color=NAVY); doc.add_paragraph(FINALTEST["cloze_context"])
for i,q in enumerate(picked,1):
    stem=q.get("stem", f'完形填空 第 {q.get("blank","")} 空')
    para(f"{i}. [{q['qid']}] {stem}")
    doc.add_paragraph("   "+"    ".join(f"{LET[j]}. {o}" for j,o in enumerate(q['options'])))
para("【答案】",bold=True,size=10,color=ORANGE)
doc.add_paragraph("  ".join(f"{i}.{LET[q['answer_index']]}" for i,q in enumerate(picked,1)))

doc.add_paragraph()
para("—— 词汇取自人教版九年级 Unit 1 真实词表;语法依据课标考点;听力/完形/阅读/写作范文及全部题目均为原创编写,未复制课本课文原文。",italic=True,size=8,color=GREY)
doc.save(DOCX)
print("JSON ->",DIR); print("DOCX saved")
